import pdf_processing
import os
from docx import Document
import image_processing
import path_processing
import cv2

def get_docx(path, args):
    # pdf -> img 
    doc = Document()
    path_processing.init_directory('temp_image')
    content_list = pdf_processing.get_content(path, args)
    for pg, content in enumerate(content_list):
        if(isinstance(content, str)):
            doc.add_paragraph(content)
        else:
            img_path = path_processing.get_f_name(par='temp_image', fname='image', ex='.png')
            cv2.imwrite(img_path, content)
            doc.add_picture(img_path)
    path_processing.clear_directory_content('temp_image')
    return doc
